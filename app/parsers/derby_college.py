from __future__ import annotations

import io
import logging
import re
from datetime import datetime

from docx import Document

from app.parsers.bases import SingleVenueParser
from app.parsers.registry import register_parser
from app.schemas import ExtractedEvent

logger = logging.getLogger(__name__)

# Date regex: "27th September", "1st Feb", "23rd Nov"
_DATE_RE = re.compile(
    r"(\d{1,2})(?:st|nd|rd|th)\s+"
    r"(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)",
    re.IGNORECASE,
)

# Date range in What's On: "25th October-2nd November"
_DATE_RANGE_RE = re.compile(
    r"(\d{1,2}(?:st|nd|rd|th)\s+\w+)\s*-\s*(\d{1,2}(?:st|nd|rd|th)\s+\w+)",
    re.IGNORECASE,
)

# Non-competition keywords to skip from What's On
_SKIP_KEYWORDS = [
    "arena hire", "clinic", "lecture", "gridwork", "course arena hire",
    "dog show", "dog agility", "reiki", "groundwork", "hobby horse",
    "he event", "ross cooper",
]


def _infer_year(month_num: int, schedule_start_year: int) -> int:
    """Infer year from month for an academic year schedule.

    Sep-Dec → start_year, Jan-Aug → start_year + 1.
    """
    return schedule_start_year if month_num >= 9 else schedule_start_year + 1


def _parse_date_no_year(text: str, schedule_start_year: int) -> str | None:
    """Parse a date like '27th September' into YYYY-MM-DD using academic year context."""
    m = _DATE_RE.search(text)
    if not m:
        return None
    day = int(m.group(1))
    month_str = m.group(2)
    # Normalise abbreviated months
    for fmt in ["%B", "%b"]:
        try:
            month_num = datetime.strptime(month_str, fmt).month
            year = _infer_year(month_num, schedule_start_year)
            return datetime(year, month_num, day).strftime("%Y-%m-%d")
        except ValueError:
            continue
    # Handle "Sept" which strptime doesn't recognise
    if month_str.lower().startswith("sep"):
        year = _infer_year(9, schedule_start_year)
        return datetime(year, 9, day).strftime("%Y-%m-%d")
    return None


@register_parser("derby_college")
class DerbyCollegeParser(SingleVenueParser):
    """Parser for Derby College Equestrian Centre — events published as .docx files.

    The main page links to WordPress Download Manager wrapper pages.
    Each wrapper embeds a data-downloadurl pointing to the actual .docx file.
    """

    VENUE_NAME = "Derby College"
    VENUE_POSTCODE = "DE7 6DN"

    async def fetch_and_parse(self, url: str) -> list[ExtractedEvent]:
        async with self._make_client() as client:
            # Step 1: Find all .docx download links from the main page
            download_links = await self._find_download_links(client, url)
            logger.info("Derby College: found %d download links", len(download_links))

            # Separate dedicated schedules from What's On
            schedule_links = []
            whatson_links = []
            for link_url, link_text in download_links:
                if "whats-on" in link_url.lower():
                    whatson_links.append((link_url, link_text))
                else:
                    schedule_links.append((link_url, link_text))

            # Process dedicated schedules first (they have class details)
            all_competitions: list[ExtractedEvent] = []
            seen_dates: set[str] = set()

            for link_url, link_text in schedule_links:
                comps = await self._process_download(client, link_url, link_text)
                for c in comps:
                    all_competitions.append(c)
                    seen_dates.add(c.date_start)

            # Process What's On, but skip dates already covered by schedules
            for link_url, link_text in whatson_links:
                comps = await self._process_download(client, link_url, link_text)
                for c in comps:
                    if c.date_start not in seen_dates:
                        all_competitions.append(c)

        self._log_result("Derby College", len(all_competitions))
        return all_competitions

    async def _process_download(self, client, link_url, link_text):
        """Download and parse a single .docx file."""
        try:
            docx_url = await self._resolve_download_url(client, link_url)
            if not docx_url:
                logger.debug("Derby College: no download URL found for %s", link_url)
                return []
            docx_bytes = await self._download_docx(client, docx_url)
            if not docx_bytes:
                return []
            comps = self._parse_docx(docx_bytes, link_url)
            logger.info("Derby College: %d events from '%s'", len(comps), link_text)
            return comps
        except Exception as e:
            logger.warning("Derby College: failed to process %s: %s", link_url, e)
            return []

    # Keywords in download URLs that indicate equestrian schedules
    _SCHEDULE_SLUGS = ["sj-", "dressage", "whats-on", "showjump", "schedule"]

    async def _find_download_links(self, client, url):
        """Find equestrian schedule download wrapper URLs from the main page."""
        soup = await self._fetch_html(client, url)

        links: list[tuple[str, str]] = []
        for a in soup.find_all("a", href=True):
            href = a["href"]
            text = a.get_text(strip=True)
            href_lower = href.lower()
            # Only match download links with equestrian-related slugs
            if "/download/" in href_lower:
                if any(slug in href_lower for slug in self._SCHEDULE_SLUGS):
                    links.append((href, text))

        return links

    async def _resolve_download_url(self, client, wrapper_url):
        """Fetch the WPDM wrapper page and extract the actual download URL."""
        soup = await self._fetch_html(client, wrapper_url)

        # Look for data-downloadurl attribute on the download link
        dl_link = soup.find("a", class_="wpdm-download-link")
        if dl_link and dl_link.get("data-downloadurl"):
            return dl_link["data-downloadurl"]

        # Fallback: look for any link with wpdmdl parameter
        for a in soup.find_all("a", href=True):
            if "wpdmdl=" in a["href"]:
                return a["href"]

        return None

    async def _download_docx(self, client, url):
        """Download the actual .docx file."""
        resp = await client.get(url)
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "")
        # WPDM serves .docx as application/zip
        if "html" in content_type and len(resp.content) < 5000:
            logger.debug("Derby College: got HTML instead of .docx from %s", url)
            return None
        return resp.content

    def _parse_docx(self, docx_bytes, source_url):
        """Parse a .docx file and extract competitions."""
        doc = Document(io.BytesIO(docx_bytes))

        # Detect document type from content
        # Normalise curly quotes to straight for matching
        all_text = " ".join(p.text for p in doc.paragraphs).replace("\u2019", "'").lower()

        if "what's on" in all_text:
            return self._parse_whats_on(doc, source_url)
        elif "show jumping" in all_text or "showjumping" in all_text:
            return self._parse_schedule(doc, source_url, "Show Jumping")
        elif "dressage" in all_text:
            return self._parse_schedule(doc, source_url, "Dressage")
        return []

    def _detect_schedule_year(self, doc, source_url=""):
        """Detect the academic year start from document text or URL."""
        all_text = " ".join(p.text for p in doc.paragraphs)
        search_text = all_text + " " + source_url

        # Look for academic year range: "25-26", "2025-26", "2025/26"
        m = re.search(r"(\d{2,4})[-/](\d{2,4})", search_text)
        if m:
            start = m.group(1)
            if len(start) == 2:
                return 2000 + int(start)
            return int(start)

        # Look for standalone year after month/season
        m = re.search(
            r"(?:sep|oct|nov|dec|jan|feb|mar|apr|may|jun|jul|aug)\w*[-\s]+(\d{4})",
            search_text,
            re.IGNORECASE,
        )
        if m:
            return int(m.group(1))

        return datetime.now().year

    def _parse_schedule(self, doc, source_url, discipline):
        """Parse SJ or Dressage schedule .docx (table with dates in col 0, classes across)."""
        schedule_year = self._detect_schedule_year(doc, source_url)
        competitions: list[ExtractedEvent] = []

        # Find the main schedule table (first table with >2 rows)
        schedule_table = None
        for table in doc.tables:
            if len(table.rows) > 2:
                schedule_table = table
                break

        if not schedule_table:
            return competitions

        # Row 0 is headers (class names)
        header_row = schedule_table.rows[0]
        class_names = [cell.text.strip() for cell in header_row.cells[1:]]

        # Rows 1+ are event dates
        for row in schedule_table.rows[1:]:
            cells = row.cells
            date_text = cells[0].text.strip()
            if not date_text:
                continue

            date_start = _parse_date_no_year(date_text, schedule_year)
            if not date_start:
                continue

            # Collect class descriptions for this date
            classes = []
            is_special = False
            for i, cell in enumerate(cells[1:]):
                cell_text = cell.text.strip()
                if cell_text and cell_text not in classes:
                    classes.append(cell_text)
                    if "separate schedule" in cell_text.lower():
                        is_special = True

            if is_special:
                event_name = classes[0].split("(")[0].strip() if classes else f"{discipline} Special"
                name = f"Derby College {event_name}"
            else:
                name = f"Derby College Unaffiliated {discipline}"
                classes = [
                    f"{class_names[i]}: {cells[i+1].text.strip()}"
                    for i in range(min(len(class_names), len(cells) - 1))
                    if cells[i+1].text.strip()
                ]

            competitions.append(self._build_event(
                name=name,
                date_start=date_start,
                discipline=discipline,
                classes=classes,
                url=source_url,
            ))

        return competitions

    def _parse_whats_on(self, doc, source_url):
        """Parse the What's On .docx (2-column table: Date | Event)."""
        schedule_year = self._detect_schedule_year(doc, source_url)
        competitions: list[ExtractedEvent] = []

        if not doc.tables:
            return competitions

        table = doc.tables[0]

        for row in table.rows[2:]:  # Skip announcement row and header row
            cells = row.cells
            if len(cells) < 2:
                continue

            date_text = cells[0].text.strip()
            event_text = cells[1].text.strip()

            if not date_text or not event_text:
                continue

            # Skip non-competition events
            event_lower = event_text.lower()
            if any(kw in event_lower for kw in _SKIP_KEYWORDS):
                continue

            # Only include "Derby College" events (their own competitions)
            if "derby college" not in event_lower:
                continue

            # Parse date (may be a range)
            range_match = _DATE_RANGE_RE.search(date_text)
            if range_match:
                date_start = _parse_date_no_year(range_match.group(1), schedule_year)
                date_end = _parse_date_no_year(range_match.group(2), schedule_year)
            else:
                date_start = _parse_date_no_year(date_text, schedule_year)
                date_end = None

            if not date_start:
                continue

            discipline = None

            # Standardise name to match schedule format for upsert dedup
            name = re.sub(r"(?i)\bshowjumping\b", "Show Jumping", event_text)
            name = name.split("-")[0].strip()  # Remove trailing notes like "- fancy dress"

            competitions.append(self._build_event(
                name=name,
                date_start=date_start,
                date_end=date_end,
                discipline=discipline,
                classes=[],
                url=source_url,
            ))

        return competitions
